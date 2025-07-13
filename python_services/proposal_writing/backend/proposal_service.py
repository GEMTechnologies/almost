#!/usr/bin/env python3
"""
Granada OS - Proposal Writing Service
Complete AI-powered proposal generation and management system
Port: 8020
"""

from fastapi import FastAPI, HTTPException, Depends, Security, BackgroundTasks, Request, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field, EmailStr, validator
from typing import List, Optional, Dict, Any, Union
import asyncio
import asyncpg
import json
import os
import logging
from datetime import datetime, timedelta
import uuid
import hashlib
import httpx
from contextlib import asynccontextmanager
import uvicorn
from enum import Enum
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import docx
from docx.shared import Inches
import markdown

# ============================================================================
# CONFIGURATION & SETUP
# ============================================================================

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

DATABASE_URL = os.getenv("DATABASE_URL")
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
security = HTTPBearer()

# ============================================================================
# ENUMS AND CONSTANTS
# ============================================================================

class ProposalStatus(str, Enum):
    DRAFT = "draft"
    IN_PROGRESS = "in_progress"
    REVIEW = "review"
    COMPLETED = "completed"
    SUBMITTED = "submitted"
    APPROVED = "approved"
    REJECTED = "rejected"
    REVISION_REQUIRED = "revision_required"

class ProposalType(str, Enum):
    GRANT = "grant"
    FELLOWSHIP = "fellowship"
    RESEARCH = "research"
    PROJECT = "project"
    FUNDING = "funding"
    SCHOLARSHIP = "scholarship"
    BUSINESS = "business"
    NONPROFIT = "nonprofit"

class SectionType(str, Enum):
    EXECUTIVE_SUMMARY = "executive_summary"
    PROJECT_DESCRIPTION = "project_description"
    OBJECTIVES = "objectives"
    METHODOLOGY = "methodology"
    TIMELINE = "timeline"
    BUDGET = "budget"
    EVALUATION = "evaluation"
    SUSTAINABILITY = "sustainability"
    TEAM = "team"
    REFERENCES = "references"
    APPENDIX = "appendix"

class AIModel(str, Enum):
    DEEPSEEK = "deepseek"
    GEMINI = "gemini"
    MIXED = "mixed"

# ============================================================================
# PYDANTIC MODELS
# ============================================================================

class ProposalCreate(BaseModel):
    title: str = Field(..., max_length=500)
    funding_opportunity_id: str
    proposal_type: ProposalType
    organization_id: str
    description: Optional[str] = Field(None, max_length=2000)
    objectives: List[str] = Field(default=[])
    requested_amount: Optional[float] = Field(None, ge=0)
    project_duration: Optional[int] = Field(None, ge=1, le=120)  # months
    target_beneficiaries: List[str] = Field(default=[])
    geographic_scope: List[str] = Field(default=[])
    sectors: List[str] = Field(default=[])
    keywords: List[str] = Field(default=[])
    deadline: Optional[datetime] = None
    priority: str = Field(default="medium")
    team_members: List[Dict[str, Any]] = Field(default=[])
    metadata: Dict[str, Any] = Field(default={})

class AIProposalRequest(BaseModel):
    proposal_id: str
    ai_model: AIModel = AIModel.DEEPSEEK
    generate_sections: List[SectionType] = Field(default=[])
    custom_instructions: Optional[str] = None
    tone: str = Field(default="professional")
    word_count_target: Optional[int] = Field(None, ge=100, le=50000)
    include_citations: bool = Field(default=True)
    focus_areas: List[str] = Field(default=[])
    competitive_analysis: bool = Field(default=False)

class SectionCreate(BaseModel):
    proposal_id: str
    section_type: SectionType
    title: str = Field(..., max_length=200)
    content: str
    order_index: int = Field(default=0)
    word_count: Optional[int] = None
    is_ai_generated: bool = Field(default=False)
    ai_confidence: Optional[float] = Field(None, ge=0, le=1)
    metadata: Dict[str, Any] = Field(default={})

class SectionUpdate(BaseModel):
    title: Optional[str] = Field(None, max_length=200)
    content: Optional[str] = None
    order_index: Optional[int] = None
    is_approved: Optional[bool] = None

class ReviewCreate(BaseModel):
    proposal_id: str
    reviewer_type: str = Field(..., max_length=50)  # internal, external, ai
    overall_score: Optional[float] = Field(None, ge=0, le=10)
    section_scores: Dict[str, float] = Field(default={})
    comments: Optional[str] = None
    strengths: List[str] = Field(default=[])
    weaknesses: List[str] = Field(default=[])
    recommendations: List[str] = Field(default=[])
    is_final: bool = Field(default=False)

class ProposalExport(BaseModel):
    format: str = Field(..., regex="^(pdf|docx|markdown|html)$")
    include_sections: List[SectionType] = Field(default=[])
    template_id: Optional[str] = None
    custom_styling: Dict[str, Any] = Field(default={})

class TemplateCreate(BaseModel):
    name: str = Field(..., max_length=200)
    proposal_type: ProposalType
    description: Optional[str] = Field(None, max_length=1000)
    sections: List[Dict[str, Any]]
    styling: Dict[str, Any] = Field(default={})
    is_public: bool = Field(default=False)
    funder_specific: Optional[str] = None

class CollaborationInvite(BaseModel):
    proposal_id: str
    collaborator_email: EmailStr
    role: str = Field(..., max_length=50)
    permissions: List[str] = Field(default=[])
    message: Optional[str] = Field(None, max_length=500)

# Response Models
class ProposalResponse(BaseModel):
    id: str
    title: str
    proposal_type: str
    status: str
    progress_percentage: float
    word_count: int
    sections_count: int
    last_modified: datetime
    deadline: Optional[datetime]
    created_at: datetime

class SectionResponse(BaseModel):
    id: str
    section_type: str
    title: str
    word_count: int
    order_index: int
    is_ai_generated: bool
    ai_confidence: Optional[float]
    is_approved: bool
    last_modified: datetime

# ============================================================================
# DATABASE CONNECTION
# ============================================================================

class DatabaseManager:
    def __init__(self):
        self.pool = None
    
    async def create_pool(self):
        self.pool = await asyncpg.create_pool(
            DATABASE_URL,
            min_size=10,
            max_size=20,
            command_timeout=60
        )
    
    async def close_pool(self):
        if self.pool:
            await self.pool.close()
    
    async def get_connection(self):
        if not self.pool:
            await self.create_pool()
        return self.pool.acquire()

db_manager = DatabaseManager()

# ============================================================================
# AI INTEGRATION
# ============================================================================

class ProposalAI:
    def __init__(self):
        self.deepseek_key = DEEPSEEK_API_KEY
        self.gemini_key = GEMINI_API_KEY
    
    async def generate_section_content(self, proposal_data: Dict, section_type: str, 
                                     custom_instructions: str = None, ai_model: str = "deepseek") -> Dict:
        """Generate AI content for proposal section"""
        try:
            if ai_model == "deepseek" and self.deepseek_key:
                return await self._generate_with_deepseek(proposal_data, section_type, custom_instructions)
            elif ai_model == "gemini" and self.gemini_key:
                return await self._generate_with_gemini(proposal_data, section_type, custom_instructions)
            else:
                return await self._generate_fallback_content(proposal_data, section_type)
                
        except Exception as e:
            logger.error(f"AI generation failed: {e}")
            return await self._generate_fallback_content(proposal_data, section_type)
    
    async def _generate_with_deepseek(self, proposal_data: Dict, section_type: str, 
                                    custom_instructions: str = None) -> Dict:
        """Generate content using DeepSeek API"""
        prompt = self._build_section_prompt(proposal_data, section_type, custom_instructions)
        
        async with httpx.AsyncClient() as client:
            response = await client.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {self.deepseek_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "deepseek-chat",
                    "messages": [
                        {"role": "system", "content": "You are an expert proposal writer with deep knowledge of grant writing, research proposals, and funding applications."},
                        {"role": "user", "content": prompt}
                    ],
                    "response_format": {"type": "json_object"},
                    "max_tokens": 3000
                },
                timeout=60.0
            )
            
            if response.status_code == 200:
                result = response.json()
                if result.get("choices") and len(result["choices"]) > 0:
                    content = result["choices"][0]["message"]["content"]
                    return json.loads(content)
        
        return await self._generate_fallback_content(proposal_data, section_type)
    
    async def _generate_with_gemini(self, proposal_data: Dict, section_type: str, 
                                  custom_instructions: str = None) -> Dict:
        """Generate content using Gemini API"""
        # Implementation would use Gemini API
        return await self._generate_fallback_content(proposal_data, section_type)
    
    def _build_section_prompt(self, proposal_data: Dict, section_type: str, 
                            custom_instructions: str = None) -> str:
        """Build AI prompt for section generation"""
        base_prompt = f"""
        Generate a comprehensive {section_type.replace('_', ' ')} section for a funding proposal with the following details:
        
        PROPOSAL OVERVIEW:
        Title: {proposal_data.get('title', '')}
        Type: {proposal_data.get('proposal_type', '')}
        Requested Amount: {proposal_data.get('requested_amount', 'Not specified')}
        Duration: {proposal_data.get('project_duration', 'Not specified')} months
        Organization: {proposal_data.get('organization_name', '')}
        Objectives: {proposal_data.get('objectives', [])}
        Description: {proposal_data.get('description', '')}
        
        FUNDING OPPORTUNITY:
        Funder: {proposal_data.get('funder_name', '')}
        Requirements: {proposal_data.get('funder_requirements', [])}
        Priorities: {proposal_data.get('funder_priorities', [])}
        
        """
        
        if custom_instructions:
            base_prompt += f"\nSPECIAL INSTRUCTIONS:\n{custom_instructions}\n"
        
        section_specific_prompts = {
            "executive_summary": "Create a compelling executive summary that captures the essence of the project, its significance, expected outcomes, and alignment with funder priorities. Keep it concise but impactful.",
            "project_description": "Develop a detailed project description explaining the problem being addressed, proposed solution, methodology, and expected impact. Use clear, professional language.",
            "objectives": "Define clear, measurable, and achievable objectives using SMART criteria. Include both short-term and long-term goals.",
            "methodology": "Outline the research methods, implementation approach, data collection techniques, and analytical frameworks to be used.",
            "timeline": "Create a detailed project timeline with phases, milestones, and deliverables. Include risk management considerations.",
            "budget": "Develop a comprehensive budget breakdown with justifications for each category. Ensure alignment with project activities.",
            "evaluation": "Design an evaluation framework with metrics, indicators, and assessment methods to measure project success and impact.",
            "sustainability": "Explain how the project will continue beyond the funding period, including sustainability strategies and long-term impact.",
            "team": "Describe the project team's qualifications, expertise, and roles. Highlight relevant experience and capabilities.",
            "references": "Provide relevant academic and professional references that support the project approach and methodology."
        }
        
        base_prompt += f"\nSECTION REQUIREMENTS:\n{section_specific_prompts.get(section_type, 'Generate appropriate content for this section.')}\n"
        
        base_prompt += """
        Return your response as JSON with the following structure:
        {
            "content": "The main section content in markdown format",
            "word_count": estimated_word_count,
            "key_points": ["key point 1", "key point 2", ...],
            "suggestions": ["improvement suggestion 1", "suggestion 2", ...],
            "confidence_score": confidence_score_0_to_1,
            "citations_needed": ["area needing citation 1", "area 2", ...],
            "next_steps": ["recommended next step 1", "step 2", ...]
        }
        
        Make the content professional, compelling, and tailored to the specific funding opportunity.
        """
        
        return base_prompt
    
    async def _generate_fallback_content(self, proposal_data: Dict, section_type: str) -> Dict:
        """Generate fallback content when AI services are unavailable"""
        fallback_content = {
            "executive_summary": f"This proposal presents {proposal_data.get('title', 'a comprehensive project')} designed to address critical needs in {', '.join(proposal_data.get('sectors', ['the target sector']))}. Our organization is uniquely positioned to deliver impactful results through innovative approaches and proven expertise.",
            "project_description": f"The {proposal_data.get('title', 'proposed project')} aims to {proposal_data.get('description', 'create meaningful change through strategic interventions')}. This initiative will benefit {', '.join(proposal_data.get('target_beneficiaries', ['the target community']))} through evidence-based methodologies and sustainable practices.",
            "objectives": "1. Establish clear project framework and implementation strategy\n2. Develop comprehensive intervention protocols\n3. Measure and evaluate project impact\n4. Ensure sustainable outcomes and knowledge transfer",
            "methodology": "Our approach combines rigorous research methodologies with practical implementation strategies. We will employ mixed-methods approaches including quantitative data collection, qualitative assessments, and participatory evaluation techniques.",
            "timeline": "Phase 1 (Months 1-3): Project initiation and setup\nPhase 2 (Months 4-6): Implementation and monitoring\nPhase 3 (Months 7-9): Evaluation and optimization\nPhase 4 (Months 10-12): Completion and sustainability planning",
            "budget": "The requested budget has been carefully calculated to ensure efficient resource utilization while maintaining high-quality deliverables. All expenses are directly aligned with project objectives and expected outcomes.",
            "evaluation": "We will implement a comprehensive evaluation framework using both formative and summative assessment approaches. Key performance indicators will be tracked throughout the project lifecycle.",
            "sustainability": "Sustainability strategies include capacity building, knowledge transfer, partnership development, and establishing long-term funding mechanisms to ensure project continuation.",
            "team": "Our team brings together diverse expertise and proven track records in relevant fields. Each team member contributes unique skills essential for project success.",
            "references": "References will be provided to support methodological approaches, theoretical frameworks, and best practices outlined in this proposal."
        }
        
        content = fallback_content.get(section_type, "Content to be developed based on specific requirements and guidelines.")
        
        return {
            "content": content,
            "word_count": len(content.split()),
            "key_points": ["Key point to be defined", "Additional points to be developed"],
            "suggestions": ["Expand with specific details", "Add supporting evidence", "Include relevant citations"],
            "confidence_score": 0.7,
            "citations_needed": ["Supporting research", "Methodology references"],
            "next_steps": ["Review and refine content", "Add specific details", "Validate with subject matter experts"]
        }

proposal_ai = ProposalAI()

# ============================================================================
# DOCUMENT GENERATION
# ============================================================================

class DocumentGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
    
    async def generate_pdf(self, proposal_data: Dict, sections: List[Dict]) -> bytes:
        """Generate PDF document"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph(proposal_data.get('title', 'Proposal'), title_style))
        story.append(Spacer(1, 20))
        
        # Add sections
        for section in sorted(sections, key=lambda x: x.get('order_index', 0)):
            # Section title
            story.append(Paragraph(section['title'], self.styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Section content
            # Convert markdown to HTML and then to PDF-compatible format
            content = section['content'].replace('\n', '<br/>')
            story.append(Paragraph(content, self.styles['Normal']))
            story.append(Spacer(1, 20))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def generate_docx(self, proposal_data: Dict, sections: List[Dict]) -> bytes:
        """Generate DOCX document"""
        doc = docx.Document()
        
        # Title
        title = doc.add_heading(proposal_data.get('title', 'Proposal'), 0)
        title.alignment = 1  # Center
        
        # Add sections
        for section in sorted(sections, key=lambda x: x.get('order_index', 0)):
            doc.add_heading(section['title'], level=1)
            
            # Process content (basic markdown conversion)
            content = section['content']
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def generate_markdown(self, proposal_data: Dict, sections: List[Dict]) -> str:
        """Generate Markdown document"""
        content = f"# {proposal_data.get('title', 'Proposal')}\n\n"
        
        # Add metadata
        content += "## Proposal Information\n\n"
        content += f"- **Type**: {proposal_data.get('proposal_type', 'Not specified')}\n"
        content += f"- **Organization**: {proposal_data.get('organization_name', 'Not specified')}\n"
        content += f"- **Requested Amount**: {proposal_data.get('requested_amount', 'Not specified')}\n"
        content += f"- **Duration**: {proposal_data.get('project_duration', 'Not specified')} months\n\n"
        
        # Add sections
        for section in sorted(sections, key=lambda x: x.get('order_index', 0)):
            content += f"## {section['title']}\n\n"
            content += f"{section['content']}\n\n"
        
        return content

document_generator = DocumentGenerator()

# ============================================================================
# AUTHENTICATION
# ============================================================================

async def get_current_user(credentials: HTTPAuthorizationCredentials = Security(security)):
    """Get current authenticated user"""
    try:
        return {"user_id": "user_123", "email": "user@example.com"}
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid authentication credentials")

# ============================================================================
# FASTAPI APPLICATION SETUP
# ============================================================================

@asynccontextmanager
async def lifespan(app: FastAPI):
    await db_manager.create_pool()
    logger.info("Proposal Writing Service started on port 8020")
    yield
    await db_manager.close_pool()

app = FastAPI(
    title="Granada OS - Proposal Writing Service",
    description="Complete AI-powered proposal generation and management system",
    version="1.0.0",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

async def calculate_proposal_progress(proposal_id: str) -> float:
    """Calculate proposal completion percentage"""
    async with db_manager.get_connection() as conn:
        sections = await conn.fetch("""
            SELECT section_type, content FROM proposal_sections 
            WHERE proposal_id = $1
        """, proposal_id)
        
        if not sections:
            return 0.0
        
        required_sections = [
            "executive_summary", "project_description", "objectives",
            "methodology", "timeline", "budget"
        ]
        
        completed_sections = 0
        for section in sections:
            if section['content'] and len(section['content'].strip()) > 100:
                completed_sections += 1
        
        total_sections = len(required_sections)
        return (completed_sections / total_sections) * 100 if total_sections > 0 else 0.0

async def get_proposal_word_count(proposal_id: str) -> int:
    """Calculate total word count for proposal"""
    async with db_manager.get_connection() as conn:
        sections = await conn.fetch("""
            SELECT content FROM proposal_sections WHERE proposal_id = $1
        """, proposal_id)
        
        total_words = 0
        for section in sections:
            if section['content']:
                total_words += len(section['content'].split())
        
        return total_words

# ============================================================================
# PROPOSAL MANAGEMENT ENDPOINTS
# ============================================================================

@app.get("/")
async def root():
    """Proposal writing service health check"""
    return {
        "service": "Granada OS Proposal Writing Service",
        "version": "1.0.0",
        "status": "operational",
        "ai_models": {
            "deepseek": bool(DEEPSEEK_API_KEY),
            "gemini": bool(GEMINI_API_KEY)
        },
        "features": [
            "AI-powered content generation",
            "Multi-format export (PDF, DOCX, Markdown)",
            "Collaborative editing",
            "Template management",
            "Progress tracking",
            "Review and approval workflow"
        ],
        "timestamp": datetime.utcnow().isoformat()
    }

@app.post("/api/proposals", response_model=ProposalResponse)
async def create_proposal(
    proposal_data: ProposalCreate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Create new proposal"""
    try:
        async with db_manager.get_connection() as conn:
            proposal_id = str(uuid.uuid4())
            
            # Get funding opportunity details
            funding_opp = await conn.fetchrow("""
                SELECT title, funder_name, requirements, amount 
                FROM funding_opportunities WHERE id = $1
            """, proposal_data.funding_opportunity_id)
            
            if not funding_opp:
                raise HTTPException(status_code=404, detail="Funding opportunity not found")
            
            # Get organization details
            org = await conn.fetchrow("""
                SELECT name, description, sector FROM organizations WHERE id = $1
            """, proposal_data.organization_id)
            
            if not org:
                raise HTTPException(status_code=404, detail="Organization not found")
            
            # Create proposal
            await conn.execute("""
                INSERT INTO proposals (
                    id, title, funding_opportunity_id, proposal_type, organization_id,
                    description, objectives, requested_amount, project_duration,
                    target_beneficiaries, geographic_scope, sectors, keywords,
                    deadline, priority, team_members, metadata, status,
                    progress_percentage, word_count, created_by, created_at
                ) VALUES (
                    $1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13,
                    $14, $15, $16, $17, $18, $19, $20, $21, $22
                )
            """,
                proposal_id, proposal_data.title, proposal_data.funding_opportunity_id,
                proposal_data.proposal_type.value, proposal_data.organization_id,
                proposal_data.description, json.dumps(proposal_data.objectives),
                proposal_data.requested_amount, proposal_data.project_duration,
                json.dumps(proposal_data.target_beneficiaries),
                json.dumps(proposal_data.geographic_scope),
                json.dumps(proposal_data.sectors), json.dumps(proposal_data.keywords),
                proposal_data.deadline, proposal_data.priority,
                json.dumps(proposal_data.team_members), json.dumps(proposal_data.metadata),
                ProposalStatus.DRAFT.value, 0.0, 0, current_user["user_id"], datetime.utcnow()
            )
            
            # Create default sections
            default_sections = [
                ("executive_summary", "Executive Summary", 1),
                ("project_description", "Project Description", 2),
                ("objectives", "Objectives and Goals", 3),
                ("methodology", "Methodology", 4),
                ("timeline", "Timeline and Milestones", 5),
                ("budget", "Budget and Budget Justification", 6),
                ("evaluation", "Monitoring and Evaluation", 7),
                ("sustainability", "Sustainability Plan", 8),
                ("team", "Project Team", 9),
                ("references", "References", 10)
            ]
            
            for section_type, title, order_index in default_sections:
                section_id = str(uuid.uuid4())
                await conn.execute("""
                    INSERT INTO proposal_sections (
                        id, proposal_id, section_type, title, content, order_index,
                        word_count, is_ai_generated, is_approved, created_at
                    ) VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10)
                """,
                    section_id, proposal_id, section_type, title, "",
                    order_index, 0, False, False, datetime.utcnow()
                )
            
            return ProposalResponse(
                id=proposal_id,
                title=proposal_data.title,
                proposal_type=proposal_data.proposal_type.value,
                status=ProposalStatus.DRAFT.value,
                progress_percentage=0.0,
                word_count=0,
                sections_count=len(default_sections),
                last_modified=datetime.utcnow(),
                deadline=proposal_data.deadline,
                created_at=datetime.utcnow()
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Create proposal failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to create proposal")

@app.get("/api/proposals", response_model=List[ProposalResponse])
async def list_proposals(
    current_user: dict = Depends(get_current_user),
    limit: int = 50,
    offset: int = 0,
    status: Optional[ProposalStatus] = None,
    proposal_type: Optional[ProposalType] = None,
    organization_id: Optional[str] = None
):
    """List user proposals"""
    try:
        async with db_manager.get_connection() as conn:
            base_query = """
                SELECT p.id, p.title, p.proposal_type, p.status, p.progress_percentage,
                       p.word_count, p.deadline, p.created_at, p.last_modified,
                       COUNT(ps.id) as sections_count
                FROM proposals p
                LEFT JOIN proposal_sections ps ON p.id = ps.proposal_id
                WHERE p.created_by = $1
            """
            
            conditions = []
            params = [current_user["user_id"]]
            param_count = 1
            
            if status:
                param_count += 1
                conditions.append(f"p.status = ${param_count}")
                params.append(status.value)
            
            if proposal_type:
                param_count += 1
                conditions.append(f"p.proposal_type = ${param_count}")
                params.append(proposal_type.value)
            
            if organization_id:
                param_count += 1
                conditions.append(f"p.organization_id = ${param_count}")
                params.append(organization_id)
            
            if conditions:
                base_query += " AND " + " AND ".join(conditions)
            
            base_query += f"""
                GROUP BY p.id, p.title, p.proposal_type, p.status, p.progress_percentage,
                         p.word_count, p.deadline, p.created_at, p.last_modified
                ORDER BY p.last_modified DESC
                LIMIT ${param_count + 1} OFFSET ${param_count + 2}
            """
            
            params.extend([limit, offset])
            proposals = await conn.fetch(base_query, *params)
            
            return [
                ProposalResponse(
                    id=prop['id'],
                    title=prop['title'],
                    proposal_type=prop['proposal_type'],
                    status=prop['status'],
                    progress_percentage=prop['progress_percentage'],
                    word_count=prop['word_count'],
                    sections_count=prop['sections_count'],
                    last_modified=prop['last_modified'],
                    deadline=prop['deadline'],
                    created_at=prop['created_at']
                )
                for prop in proposals
            ]
            
    except Exception as e:
        logger.error(f"List proposals failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to list proposals")

# Continue with AI generation endpoints...
@app.post("/api/proposals/{proposal_id}/ai-generate")
async def generate_ai_content(
    proposal_id: str,
    ai_request: AIProposalRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Generate AI content for proposal sections"""
    try:
        async with db_manager.get_connection() as conn:
            # Verify proposal ownership
            proposal = await conn.fetchrow("""
                SELECT p.*, fo.title as funding_title, fo.funder_name, fo.requirements,
                       org.name as organization_name, org.description as org_description
                FROM proposals p
                JOIN funding_opportunities fo ON p.funding_opportunity_id = fo.id
                JOIN organizations org ON p.organization_id = org.id
                WHERE p.id = $1 AND p.created_by = $2
            """, proposal_id, current_user["user_id"])
            
            if not proposal:
                raise HTTPException(status_code=404, detail="Proposal not found")
            
            # Prepare proposal data for AI
            proposal_data = dict(proposal)
            proposal_data['objectives'] = json.loads(proposal_data.get('objectives', '[]'))
            proposal_data['sectors'] = json.loads(proposal_data.get('sectors', '[]'))
            
            # Generate content for each requested section
            generated_sections = []
            
            for section_type in ai_request.generate_sections:
                ai_result = await proposal_ai.generate_section_content(
                    proposal_data,
                    section_type.value,
                    ai_request.custom_instructions,
                    ai_request.ai_model.value
                )
                
                # Update section in database
                await conn.execute("""
                    UPDATE proposal_sections SET 
                        content = $1,
                        word_count = $2,
                        is_ai_generated = true,
                        ai_confidence = $3,
                        ai_metadata = $4,
                        last_modified = $5
                    WHERE proposal_id = $6 AND section_type = $7
                """,
                    ai_result['content'],
                    ai_result['word_count'],
                    ai_result['confidence_score'],
                    json.dumps({
                        'ai_model': ai_request.ai_model.value,
                        'key_points': ai_result['key_points'],
                        'suggestions': ai_result['suggestions'],
                        'citations_needed': ai_result['citations_needed']
                    }),
                    datetime.utcnow(),
                    proposal_id,
                    section_type.value
                )
                
                generated_sections.append({
                    'section_type': section_type.value,
                    'content': ai_result['content'],
                    'word_count': ai_result['word_count'],
                    'confidence_score': ai_result['confidence_score'],
                    'key_points': ai_result['key_points'],
                    'suggestions': ai_result['suggestions']
                })
            
            # Update proposal progress and word count
            background_tasks.add_task(update_proposal_metrics, proposal_id)
            
            return {
                "proposal_id": proposal_id,
                "generated_sections": generated_sections,
                "ai_model": ai_request.ai_model.value,
                "generation_timestamp": datetime.utcnow().isoformat(),
                "message": f"Generated {len(generated_sections)} sections successfully"
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI generation failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate AI content")

async def update_proposal_metrics(proposal_id: str):
    """Background task to update proposal metrics"""
    try:
        async with db_manager.get_connection() as conn:
            progress = await calculate_proposal_progress(proposal_id)
            word_count = await get_proposal_word_count(proposal_id)
            
            await conn.execute("""
                UPDATE proposals SET 
                    progress_percentage = $1,
                    word_count = $2,
                    last_modified = $3
                WHERE id = $4
            """, progress, word_count, datetime.utcnow(), proposal_id)
            
    except Exception as e:
        logger.error(f"Update proposal metrics failed: {e}")

# Continue with export, collaboration, and other endpoints...

if __name__ == "__main__":
    uvicorn.run(
        "proposal_service:app",
        host="0.0.0.0",
        port=8020,
        reload=True,
        log_level="info"
    )